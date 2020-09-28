import docx
import os
os.chdir(r'C:\Users\SOLO\Documents\PythonFiles')

guestsFile = open('guests.txt')
guests = []
for guest in guestsFile.readlines():
    guests.append(guest.strip())


doc = docx.Document()

for guest in guests:
    para1 = doc.add_paragraph('It would be a pleasure to have the company of')
    para2 = doc.add_paragraph(guest)
    para3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    para4 = doc.add_paragraph('April 1st')
    para5 = doc.add_paragraph('at 7 o\'clock')
    para1.runs[0].italic = True
    para2.runs[0].bold = True
    para3.runs[0].italic = True
    para5.runs[0].italic = True
    doc.add_page_break()

doc.save('invitations.docx')
    
