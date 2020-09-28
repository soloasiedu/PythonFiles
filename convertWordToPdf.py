import win32com.client
import docx
import os
os.chdir(r'C:\Users\SOLO\Documents\PythonFiles')

wordFilename = 'refund.docx'
pdfFilename = 'refund.pdf'

doc = docx.Document()

docToBeCopied = docx.Document(wordFilename)

for i in range(len(docToBeCopied.paragraphs)):
    doc.add_paragraph(docToBeCopied.paragraphs[i].text)

doc.save(wordFilename)
wdFormatPDF = 17
wordObj = win32com.client.Dispatch('Word.Application')
docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()
